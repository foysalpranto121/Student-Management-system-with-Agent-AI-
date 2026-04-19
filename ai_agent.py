"""
AI Agent Module for Student Management System
LangChain + RAG + OpenAI + ChromaDB
Features: Student insights, grade analysis, attendance tracking
"""

import os
import sqlite3
import json
from datetime import datetime
from typing import List, Dict, Any
from dotenv import load_dotenv

# Load environment variables
env_path = os.path.join(os.path.dirname(__file__), '.env')
print(f"Loading .env from: {env_path}")
loaded = load_dotenv(env_path)
print(f".env loaded successfully: {loaded}")

# LangChain imports
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_community.vectorstores import Chroma
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain.chains import ConversationalRetrievalChain, LLMChain, RetrievalQA
from langchain.memory import ConversationBufferMemory, ConversationSummaryMemory
from langchain.prompts import PromptTemplate, ChatPromptTemplate
from langchain.schema import Document, HumanMessage, AIMessage
from langchain.agents import Tool, AgentExecutor, create_react_agent
from langchain.tools import StructuredTool
from langchain.retrievers import ContextualCompressionRetriever
# from langchain.retrievers.document_compressors import LLMChainExtractor
from langchain.evaluation import load_evaluator
import numpy as np
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity

# Multi-modal imports
import base64
import io
from PIL import Image
import pytesseract
import fitz  # PyMuPDF for PDF processing
import docx
from pathlib import Path
import mimetypes
import hashlib

# Initialize OpenAI API key from environment
OPENAI_API_KEY = os.getenv('OPENAI_API_KEY')
print(f"API Key loaded: {OPENAI_API_KEY[:20] + '...' if OPENAI_API_KEY else 'NOT FOUND'}")
if not OPENAI_API_KEY or OPENAI_API_KEY == 'your-openai-api-key-here':
    print("WARNING: OpenAI API key not found or still set to default in .env")
    OPENAI_API_KEY = None
else:
    print("SUCCESS: OpenAI API key loaded successfully")

class SMSAIAgent:
    """
    Enhanced AI Agent for Student Management System
    Features: Advanced RAG, semantic search, predictive analytics, multi-modal support
    """
    
    def __init__(self, db_path='student_management.db'):
        self.db_path = db_path
        self.llm = None
        self.vector_store = None
        self.compression_retriever = None
        self.memory = None
        self.summary_memory = None
        self.qa_chain = None
        self.agent_executor = None
        self.tools = []
        self.conversation_history = []
        self.session_context = {}
        self.embedding_cache = {}
        self.document_store = {}  # Store processed documents
        self.uploaded_files = {}   # Track uploaded files
        self.cache = {}            # In-memory cache for frequently accessed data
        self.cache_timestamps = {} # Track cache freshness
        self.sync_queue = []       # Queue for data synchronization
        self.last_sync_time = None # Track last synchronization time
        
        # Initialize components with better error handling
        if OPENAI_API_KEY:
            try:
                self._init_llm()
                self._init_memory()
                self._init_vector_store()
                self._init_compression_retriever()
                self._init_tools()
                self._init_agent()
            except Exception as e:
                print(f"Error during AI component initialization: {e}")
        else:
            print("AI agent initialization skipped: No valid OpenAI API key provided.")
    
    def _init_llm(self):
        """Initialize OpenAI LLM with error handling"""
        try:
            if not OPENAI_API_KEY:
                print("Cannot initialize LLM: API key not available")
                self.llm = None
                return
                
            # Use newer API format compatible with latest openai package
            self.llm = ChatOpenAI(
                temperature=0.7,
                model='gpt-3.5-turbo',
                api_key=OPENAI_API_KEY
            )
            print("SUCCESS: LLM initialized successfully")
        except Exception as e:
            print(f"Error initializing LLM: {e}")
            import traceback
            traceback.print_exc()
            self.llm = None
    
    def _init_memory(self):
        """Initialize enhanced conversation memory"""
        self.memory = ConversationBufferMemory(
            memory_key='chat_history',
            return_messages=True,
            output_key='answer'
        )
        
        # Add summary memory for long conversations (only if LLM is available)
        if self.llm:
            try:
                self.summary_memory = ConversationSummaryMemory(
                    llm=self.llm,
                    memory_key='summary',
                    return_messages=True
                )
            except Exception as e:
                print(f"Warning: Could not initialize summary memory: {e}")
                self.summary_memory = None
        else:
            self.summary_memory = None
    
    def _get_db_connection(self):
        """Get SQLite database connection"""
        conn = sqlite3.connect(self.db_path)
        conn.row_factory = sqlite3.Row
        return conn
    
    def _fetch_student_data(self) -> List[Document]:
        """Enhanced fetch with comprehensive student data and metadata"""
        conn = self._get_db_connection()
        cursor = conn.cursor()
        documents = []
        
        try:
            # Fetch students with enhanced metadata
            cursor.execute('''
                SELECT id, student_id, name, email, department, semester, 
                       date_of_birth, address, phone, created_at, updated_at
                FROM students
                ORDER BY department, semester, name
            ''')
            students = cursor.fetchall()
            
            for student in students:
                # Enhanced student profile with temporal data
                content = f"""
                Student Profile: {student['name']} (ID: {student['student_id']})
                Department: {student['department']} | Semester: {student['semester']}
                Contact: {student['email']} | {student['phone']}
                Address: {student['address']}
                Enrolled: {student['created_at']}
                """
                
                # Academic performance with trends
                cursor.execute('''
                    SELECT subject, marks, max_marks, exam_type, grade, created_at
                    FROM grades WHERE student_id = ?
                    ORDER BY created_at DESC
                ''', (student['id'],))
                grades = cursor.fetchall()
                
                if grades:
                    content += "\n\nAcademic Performance:\n"
                    total_marks = 0
                    total_max = 0
                    subject_trends = {}
                    
                    for grade in grades:
                        content += f"  - {grade['subject']}: {grade['marks']}/{grade['max_marks']} ({grade['grade']}) [{grade['exam_type']}]\n"
                        total_marks += grade['marks']
                        total_max += grade['max_marks']
                        
                        # Track subject trends
                        if grade['subject'] not in subject_trends:
                            subject_trends[grade['subject']] = []
                        subject_trends[grade['subject']].append(grade['marks'] / grade['max_marks'] * 100)
                    
                    if total_max > 0:
                        avg_percentage = (total_marks / total_max) * 100
                        content += f"\nOverall Average: {avg_percentage:.1f}%\n"
                    
                    # Add trend analysis
                    content += "\nPerformance Trends:\n"
                    for subject, scores in subject_trends.items():
                        if len(scores) > 1:
                            trend = "improving" if scores[-1] > scores[0] else "declining"
                            content += f"  - {subject}: {trend} ({scores[0]:.1f}% → {scores[-1]:.1f}%)\n"
                
                # Enhanced attendance with patterns
                cursor.execute('''
                    SELECT date, status, created_at
                    FROM attendance WHERE student_id = ?
                    ORDER BY date DESC
                ''', (student['id'],))
                attendance_records = cursor.fetchall()
                
                if attendance_records:
                    total = len(attendance_records)
                    present = sum(1 for r in attendance_records if r['status'] == 'present')
                    attendance_pct = (present / total) * 100 if total > 0 else 0
                    content += f"\nAttendance: {attendance_pct:.1f}% ({present}/{total} days)\n"
                    
                    # Analyze attendance patterns
                    recent_attendance = attendance_records[:10]  # Last 10 records
                    recent_present = sum(1 for r in recent_attendance if r['status'] == 'present')
                    recent_pct = (recent_present / len(recent_attendance)) * 100 if recent_attendance else 0
                    
                    if recent_pct < attendance_pct - 10:
                        content += f"⚠️ Recent attendance decline: {recent_pct:.1f}% (last 10 days)\n"
                    elif recent_pct > attendance_pct + 10:
                        content += f"📈 Recent attendance improvement: {recent_pct:.1f}% (last 10 days)\n"
                
                # Create enhanced document with rich metadata
                documents.append(Document(
                    page_content=content.strip(),
                    metadata={
                        'student_id': student['student_id'],
                        'name': student['name'],
                        'department': student['department'],
                        'semester': student['semester'],
                        'email': student['email'],
                        'type': 'student_profile',
                        'avg_grade': avg_percentage if grades else None,
                        'attendance_rate': attendance_pct if attendance_records else None,
                        'total_subjects': len(set(g['subject'] for g in grades)) if grades else 0,
                        'enrollment_date': student['created_at'],
                        'last_updated': datetime.now().isoformat()
                    }
                ))
            
            # Enhanced department analytics with predictive insights
            cursor.execute('''
                SELECT department, COUNT(*) as count, AVG(semester) as avg_sem,
                       MIN(created_at) as oldest_enrollment, MAX(created_at) as newest_enrollment
                FROM students GROUP BY department
            ''')
            depts = cursor.fetchall()
            
            for dept in depts:
                content = f"""
                Department Analytics: {dept['department']}
                Total Students: {dept['count']}
                Average Semester: {dept['avg_sem']:.1f}
                Enrollment Period: {dept['oldest_enrollment']} to {dept['newest_enrollment']}
                """
                
                # Advanced grade statistics with distribution
                cursor.execute('''
                    SELECT g.grade, COUNT(*) as count, AVG(g.marks * 100.0 / g.max_marks) as avg_pct
                    FROM grades g
                    JOIN students s ON g.student_id = s.id
                    WHERE s.department = ?
                    GROUP BY g.grade
                    ORDER BY avg_pct DESC
                ''', (dept['department'],))
                grade_stats = cursor.fetchall()
                
                if grade_stats:
                    content += "\n\nGrade Distribution & Performance:\n"
                    for gs in grade_stats:
                        content += f"  - Grade {gs['grade']}: {gs['count']} students (avg: {gs['avg_pct']:.1f}%)\n"
                
                # Subject performance analysis
                cursor.execute('''
                    SELECT g.subject, AVG(g.marks * 100.0 / g.max_marks) as avg_pct,
                           COUNT(*) as attempt_count, MIN(g.marks * 100.0 / g.max_marks) as min_pct,
                           MAX(g.marks * 100.0 / g.max_marks) as max_pct
                    FROM grades g
                    JOIN students s ON g.student_id = s.id
                    WHERE s.department = ?
                    GROUP BY g.subject
                    ORDER BY avg_pct DESC
                ''', (dept['department'],))
                subject_stats = cursor.fetchall()
                
                if subject_stats:
                    content += "\nSubject Performance Analysis:\n"
                    for ss in subject_stats:
                        content += f"  - {ss['subject']}: {ss['avg_pct']:.1f}% avg (range: {ss['min_pct']:.1f}-{ss['max_pct']:.1f}%) [{ss['attempt_count']} attempts]\n"
                
                documents.append(Document(
                    page_content=content.strip(),
                    metadata={
                        'department': dept['department'],
                        'type': 'department_analytics',
                        'student_count': dept['count'],
                        'avg_semester': dept['avg_sem'],
                        'grade_distribution': {gs['grade']: gs['count'] for gs in grade_stats} if grade_stats else {},
                        'subject_performance': {ss['subject']: ss['avg_pct'] for ss in subject_stats} if subject_stats else {},
                        'analysis_date': datetime.now().isoformat()
                    }
                ))
            
            # Add school-wide insights and trends
            cursor.execute('''
                SELECT COUNT(*) as total_students,
                       AVG(CASE WHEN g.marks IS NOT NULL THEN g.marks * 100.0 / g.max_marks END) as school_avg,
                       COUNT(DISTINCT department) as dept_count
                FROM students s
                LEFT JOIN grades g ON s.id = g.student_id
            ''')
            school_stats = cursor.fetchone()
            
            if school_stats['total_students'] > 0:
                content = f"""
                School-Wide Analytics Overview
                ====================================
                Total Enrollment: {school_stats['total_students']} students
                Departments: {school_stats['dept_count']}
                School Average Performance: {school_stats['school_avg'] or 0:.1f}%
                Analysis Date: {datetime.now().strftime('%Y-%m-%d %H:%M')}
                """
                
                documents.append(Document(
                    page_content=content.strip(),
                    metadata={
                        'type': 'school_analytics',
                        'total_students': school_stats['total_students'],
                        'department_count': school_stats['dept_count'],
                        'school_average': school_stats['school_avg'] or 0,
                        'analysis_date': datetime.now().isoformat()
                    }
                ))
            
        except Exception as e:
            print(f"Error fetching student data: {e}")
        finally:
            conn.close()
        
        return documents
    
    def _init_vector_store(self):
        """Enhanced ChromaDB vector store with semantic chunking"""
        try:
            # Fetch enhanced data from database
            documents = self._fetch_student_data()
            
            if not documents:
                # Create placeholder documents if no data
                documents = [Document(
                    page_content="Student Management System. Use available tools to query student data.",
                    metadata={'type': 'system', 'created_at': datetime.now().isoformat()}
                )]
            
            # Process student data into chunks
            print("Processing student data into chunks...")
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=100,
                length_function=len,
                separators=["\n\n", "\n", ". ", " ", ""]
            )
            
            texts = text_splitter.split_documents(documents)
            print(f"Split documents into {len(texts)} chunks.")
            
            # Enhance metadata with chunk information
            for i, text in enumerate(texts):
                text.metadata.update({
                    'chunk_id': f"chunk_{i}",
                    'chunk_index': i,
                    'total_chunks': len(texts),
                    'processing_date': datetime.now().isoformat()
                })
            
            # Create embeddings and vector store
            print("Initializing vector store (this may take a few seconds)...")
            embeddings = OpenAIEmbeddings(
                api_key=OPENAI_API_KEY,
                model="text-embedding-ada-002"
            )
            
            self.vector_store = Chroma.from_documents(
                documents=texts,
                embedding=embeddings,
                persist_directory='./chroma_db'
            )
            self.vector_store.persist()
            
            print(f"Vector store initialized with {len(texts)} chunks")
            
        except Exception as e:
            print(f"Error initializing vector store: {e}")
            self.vector_store = None
    
    def _init_compression_retriever(self):
        """Initialize contextual compression for better retrieval"""
        try:
            if self.vector_store:
                print("Setting up base retriever...")
                # Simple similarity retriever
                self.compression_retriever = self.vector_store.as_retriever(
                    search_type="similarity",
                    search_kwargs={"k": 5}
                )
        except Exception as e:
            print(f"Error initializing compression retriever: {e}")
            self.compression_retriever = None
    
    def _tool_get_student_info(self, student_name: str = None, student_id: str = None) -> str:
        """Tool: Get detailed information about a specific student"""
        conn = self._get_db_connection()
        cursor = conn.cursor()
        
        try:
            if student_id:
                cursor.execute('''
                    SELECT * FROM students WHERE student_id = ?
                ''', (student_id,))
            elif student_name:
                cursor.execute('''
                    SELECT * FROM students WHERE name LIKE ?
                ''', (f'%{student_name}%',))
            else:
                return "Please provide either student name or student ID"
            
            student = cursor.fetchone()
            if not student:
                return f"Student not found"
            
            result = f"""
            Student Information:
            Name: {student['name']}
            ID: {student['student_id']}
            Department: {student['department']}
            Semester: {student['semester']}
            Email: {student['email']}
            Phone: {student['phone']}
            """
            
            # Get grades
            cursor.execute('''
                SELECT subject, marks, max_marks, grade FROM grades WHERE student_id = ?
            ''', (student['id'],))
            grades = cursor.fetchall()
            
            if grades:
                result += "\nGrades:\n"
                for g in grades:
                    result += f"  {g['subject']}: {g['marks']}/{g['max_marks']} ({g['grade']})\n"
            
            # Get attendance
            cursor.execute('''
                SELECT COUNT(*) as total, 
                       SUM(CASE WHEN status='present' THEN 1 ELSE 0 END) as present
                FROM attendance WHERE student_id = ?
            ''', (student['id'],))
            att = cursor.fetchone()
            
            if att and att['total'] > 0:
                pct = (att['present'] / att['total']) * 100
                result += f"\nAttendance: {pct:.1f}% ({att['present']}/{att['total']})\n"
            
            return result
            
        except Exception as e:
            return f"Error: {str(e)}"
        finally:
            conn.close()
    
    def _tool_get_class_statistics(self, department: str = None) -> str:
        """Tool: Get statistics for a department or entire class"""
        conn = self._get_db_connection()
        cursor = conn.cursor()
        
        try:
            if department:
                cursor.execute('SELECT COUNT(*) as count FROM students WHERE department = ?', (department,))
                dept_count = cursor.fetchone()['count']
                
                result = f"Department: {department}\nTotal Students: {dept_count}\n\n"
                
                # Grade distribution
                cursor.execute('''
                    SELECT g.grade, COUNT(*) as count
                    FROM grades g
                    JOIN students s ON g.student_id = s.id
                    WHERE s.department = ?
                    GROUP BY g.grade
                    ORDER BY count DESC
                ''', (department,))
                grades = cursor.fetchall()
                
                if grades:
                    result += "Grade Distribution:\n"
                    for g in grades:
                        result += f"  Grade {g['grade']}: {g['count']}\n"
                
                # Average marks
                cursor.execute('''
                    SELECT AVG(g.marks * 100.0 / g.max_marks) as avg_pct
                    FROM grades g
                    JOIN students s ON g.student_id = s.id
                    WHERE s.department = ?
                ''', (department,))
                avg = cursor.fetchone()
                if avg and avg['avg_pct']:
                    result += f"\nAverage Score: {avg['avg_pct']:.1f}%\n"
            else:
                cursor.execute('SELECT COUNT(*) as count FROM students')
                total = cursor.fetchone()['count']
                
                result = f"School Statistics:\nTotal Students: {total}\n\n"
                
                # Department breakdown
                cursor.execute('''
                    SELECT department, COUNT(*) as count 
                    FROM students GROUP BY department
                ''')
                depts = cursor.fetchall()
                
                result += "Students by Department:\n"
                for d in depts:
                    result += f"  {d['department']}: {d['count']}\n"
            
            return result
            
        except Exception as e:
            return f"Error: {str(e)}"
        finally:
            conn.close()
    
    def _tool_get_attendance_summary(self, date: str = None) -> str:
        """Tool: Get attendance summary for a specific date or overall"""
        conn = self._get_db_connection()
        cursor = conn.cursor()
        
        try:
            if date:
                cursor.execute('''
                    SELECT COUNT(*) as total,
                           SUM(CASE WHEN status='present' THEN 1 ELSE 0 END) as present
                    FROM attendance WHERE date = ?
                ''', (date,))
                stats = cursor.fetchone()
                
                if stats and stats['total'] > 0:
                    pct = (stats['present'] / stats['total']) * 100
                    return f"Attendance on {date}: {pct:.1f}% ({stats['present']}/{stats['total']} present)"
                return f"No attendance records for {date}"
            else:
                cursor.execute('''
                    SELECT COUNT(DISTINCT date) as days,
                           COUNT(*) as total_records,
                           SUM(CASE WHEN status='present' THEN 1 ELSE 0 END) as present
                    FROM attendance
                ''')
                stats = cursor.fetchone()
                
                if stats and stats['total_records'] > 0:
                    pct = (stats['present'] / stats['total_records']) * 100
                    return f"Overall Attendance: {pct:.1f}%\nRecorded Days: {stats['days']}\nTotal Records: {stats['total_records']}"
                return "No attendance records found"
                
        except Exception as e:
            return f"Error: {str(e)}"
        finally:
            conn.close()
    
    def _tool_analyze_student_performance(self, student_id: str) -> str:
        """Tool: Analyze a student's academic performance and provide insights"""
        conn = self._get_db_connection()
        cursor = conn.cursor()
        
        try:
            cursor.execute('SELECT * FROM students WHERE student_id = ?', (student_id,))
            student = cursor.fetchone()
            
            if not student:
                return f"Student {student_id} not found"
            
            # Get all grades
            cursor.execute('''
                SELECT subject, marks, max_marks, grade FROM grades WHERE student_id = ?
            ''', (student['id'],))
            grades = cursor.fetchall()
            
            if not grades:
                return f"No grade records found for {student['name']}"
            
            # Calculate statistics
            subjects = []
            scores = []
            grade_counts = {'A': 0, 'B': 0, 'C': 0, 'D': 0, 'F': 0}
            
            for g in grades:
                subjects.append(g['subject'])
                pct = (g['marks'] / g['max_marks']) * 100
                scores.append(pct)
                if g['grade'] in grade_counts:
                    grade_counts[g['grade']] += 1
            
            avg_score = sum(scores) / len(scores) if scores else 0
            min_score = min(scores) if scores else 0
            max_score = max(scores) if scores else 0
            
            # Find best and worst subjects
            best_idx = scores.index(max_score)
            worst_idx = scores.index(min_score)
            
            result = f"""
Performance Analysis for {student['name']} ({student_id}):

Overall Statistics:
- Average Score: {avg_score:.1f}%
- Best Subject: {subjects[best_idx]} ({max_score:.1f}%)
- Needs Improvement: {subjects[worst_idx]} ({min_score:.1f}%)
- Total Subjects: {len(subjects)}

Grade Distribution:
- A: {grade_counts['A']}
- B: {grade_counts['B']}
- C: {grade_counts['C']}
- D: {grade_counts['D']}
- F: {grade_counts['F']}
"""
            
            # Add attendance context
            cursor.execute('''
                SELECT COUNT(*) as total, 
                       SUM(CASE WHEN status='present' THEN 1 ELSE 0 END) as present
                FROM attendance WHERE student_id = ?
            ''', (student['id'],))
            att = cursor.fetchone()
            
            if att and att['total'] > 0:
                att_pct = (att['present'] / att['total']) * 100
                result += f"\nAttendance Rate: {att_pct:.1f}%"
                if att_pct < 75:
                    result += " (⚠️ Low attendance - may affect performance)"
            
            return result
            
        except Exception as e:
            return f"Error: {str(e)}"
        finally:
            conn.close()
    
    def _tool_predict_student_performance(self, student_id: str) -> str:
        """Tool: Predict student performance using historical data trends"""
        conn = self._get_db_connection()
        cursor = conn.cursor()
        
        try:
            cursor.execute('SELECT * FROM students WHERE student_id = ?', (student_id,))
            student = cursor.fetchone()
            
            if not student:
                return f"Student {student_id} not found"
            
            # Get historical grades with temporal data
            cursor.execute('''
                SELECT subject, marks, max_marks, created_at
                FROM grades WHERE student_id = ?
                ORDER BY created_at ASC
            ''', (student['id'],))
            grades = cursor.fetchall()
            
            if len(grades) < 3:
                return f"Insufficient data for prediction (need at least 3 grade records)"
            
            # Analyze trends by subject
            subject_trends = {}
            for grade in grades:
                subject = grade['subject']
                percentage = (grade['marks'] / grade['max_marks']) * 100
                
                if subject not in subject_trends:
                    subject_trends[subject] = []
                subject_trends[subject].append({
                    'percentage': percentage,
                    'date': grade['created_at']
                })
            
            predictions = []
            overall_trend = []
            
            for subject, data in subject_trends.items():
                if len(data) >= 2:
                    # Simple linear trend prediction
                    percentages = [d['percentage'] for d in data]
                    recent_avg = sum(percentages[-3:]) / len(percentages[-3:])
                    
                    # Calculate trend direction
                    if len(percentages) >= 3:
                        early_avg = sum(percentages[:len(percentages)//2]) / (len(percentages)//2)
                        late_avg = sum(percentages[len(percentages)//2:]) / (len(percentages) - len(percentages)//2)
                        trend_direction = "improving" if late_avg > early_avg else "declining" if late_avg < early_avg else "stable"
                    else:
                        trend_direction = "stable"
                    
                    # Predict next performance
                    if trend_direction == "improving":
                        predicted = min(recent_avg + 5, 95)  # Cap at 95%
                    elif trend_direction == "declining":
                        predicted = max(recent_avg - 5, 35)  # Floor at 35%
                    else:
                        predicted = recent_avg
                    
                    predictions.append({
                        'subject': subject,
                        'current_avg': recent_avg,
                        'predicted': predicted,
                        'trend': trend_direction,
                        'confidence': 'high' if len(data) >= 5 else 'medium'
                    })
                    
                    overall_trend.append(predicted)
            
            # Generate prediction summary
            result = f"""
🔮 Performance Prediction for {student['name']} ({student_id}):

Predicted Performance by Subject:
"""
            
            for pred in predictions:
                trend_icon = "📈" if pred['trend'] == "improving" else "📉" if pred['trend'] == "declining" else "➡️"
                confidence_icon = "🔥" if pred['confidence'] == 'high' else "⚡"
                result += f"  {trend_icon} {pred['subject']}: {pred['current_avg']:.1f}% → {pred['predicted']:.1f}% {confidence_icon}\n"
            
            if overall_trend:
                overall_predicted = sum(overall_trend) / len(overall_trend)
                result += f"\n📊 Overall Predicted Performance: {overall_predicted:.1f}%\n"
                
                # Add recommendations
                if overall_predicted < 60:
                    result += "\n⚠️ Recommendations: Consider additional tutoring and study support\n"
                elif overall_predicted > 85:
                    result += "\n🌟 Recommendations: Student shows excellent potential, consider advanced challenges\n"
                else:
                    result += "\n✅ Recommendations: Continue current study approach with minor adjustments\n"
            
            return result.strip()
            
        except Exception as e:
            return f"Error generating prediction: {str(e)}"
        finally:
            conn.close()
    
    def _tool_analyze_attendance_patterns(self, days: int = 30) -> str:
        """Tool: Analyze attendance patterns and identify at-risk students"""
        conn = self._get_db_connection()
        cursor = conn.cursor()
        
        try:
            # Get recent attendance data
            cursor.execute('''
                SELECT s.student_id, s.name, s.department,
                       COUNT(*) as total_days,
                       SUM(CASE WHEN a.status = 'present' THEN 1 ELSE 0 END) as present_days,
                       AVG(CASE WHEN a.status = 'present' THEN 1 ELSE 0 END) as attendance_rate
                FROM students s
                LEFT JOIN attendance a ON s.id = a.student_id 
                    AND a.date >= date('now', '-{} days')
                GROUP BY s.id
                HAVING total_days > 0
                ORDER BY attendance_rate ASC
            '''.format(days))
            
            attendance_data = cursor.fetchall()
            
            if not attendance_data:
                return f"No attendance data found for the last {days} days"
            
            result = f"📅 Attendance Pattern Analysis (Last {days} days):\n\n"
            
            # Identify at-risk students (attendance < 75%)
            at_risk = [row for row in attendance_data if row['attendance_rate'] < 0.75]
            good_attendance = [row for row in attendance_data if row['attendance_rate'] >= 0.90]
            
            result += f"📊 Summary:\n"
            result += f"  - Total Students with Records: {len(attendance_data)}\n"
            result += f"  - At-Risk Students (<75%): {len(at_risk)}\n"
            result += f"  - Excellent Attendance (≥90%): {len(good_attendance)}\n\n"
            
            if at_risk:
                result += "⚠️ At-Risk Students:\n"
                for student in at_risk[:5]:  # Show top 5 at-risk
                    rate = student['attendance_rate'] * 100
                    result += f"  - {student['name']} ({student['student_id']}): {rate:.1f}% - {student['department']}\n"
                
                if len(at_risk) > 5:
                    result += f"  ... and {len(at_risk) - 5} more students\n"
                
                result += "\n💡 Recommendations for At-Risk Students:\n"
                result += "  1. Schedule parent-teacher meetings\n"
                result += "  2. Implement attendance improvement plans\n"
                result += "  3. Provide academic support if needed\n"
            
            if good_attendance:
                result += f"\n🌟 Excellent Attendance Students: {len(good_attendance)}\n"
            
            # Department-wise analysis
            cursor.execute('''
                SELECT s.department,
                       AVG(CASE WHEN a.status = 'present' THEN 1 ELSE 0 END) as avg_attendance,
                       COUNT(DISTINCT s.id) as student_count
                FROM students s
                LEFT JOIN attendance a ON s.id = a.student_id 
                    AND a.date >= date('now', '-{} days')
                GROUP BY s.department
                ORDER BY avg_attendance DESC
            '''.format(days))
            
            dept_stats = cursor.fetchall()
            if dept_stats:
                result += "\n🏢 Department-wise Attendance:\n"
                for dept in dept_stats:
                    rate = dept['avg_attendance'] * 100 if dept['avg_attendance'] else 0
                    result += f"  - {dept['department']}: {rate:.1f}% ({dept['student_count']} students)\n"
            
            return result.strip()
            
        except Exception as e:
            return f"Error analyzing attendance patterns: {str(e)}"
        finally:
            conn.close()
    
    def _tool_identify_learning_trends(self, department: str = None) -> str:
        """Tool: Identify learning trends and subject performance patterns"""
        conn = self._get_db_connection()
        cursor = conn.cursor()
        
        try:
            if department:
                cursor.execute('SELECT id, name FROM students WHERE department = ?', (department,))
                students = cursor.fetchall()
                if not students:
                    return f"No students found in department: {department}"
                student_ids = [s['id'] for s in students]
                dept_filter = f"AND s.id IN ({','.join(['?']*len(student_ids))})"
                params = student_ids
            else:
                dept_filter = ""
                params = []
            
            # Get subject performance trends over time
            cursor.execute(f'''
                SELECT g.subject, 
                       AVG(g.marks * 100.0 / g.max_marks) as avg_performance,
                       COUNT(*) as attempt_count,
                       MIN(g.marks * 100.0 / g.max_marks) as min_performance,
                       MAX(g.marks * 100.0 / g.max_marks) as max_performance,
                       s.department
                FROM grades g
                JOIN students s ON g.student_id = s.id
                WHERE 1=1 {dept_filter}
                GROUP BY g.subject, s.department
                ORDER BY avg_performance DESC
            ''', params)
            
            subject_data = cursor.fetchall()
            
            if not subject_data:
                return "No grade data available for analysis"
            
            result = f"📈 Learning Trends Analysis\n"
            if department:
                result += f"Department: {department}\n"
            result += "=" * 40 + "\n\n"
            
            # Identify best and worst performing subjects
            best_subjects = subject_data[:3]
            worst_subjects = subject_data[-3:]
            
            result += "🏆 Top Performing Subjects:\n"
            for subject in best_subjects:
                result += f"  - {subject['subject']}: {subject['avg_performance']:.1f}% avg ({subject['attempt_count']} attempts)\n"
                if subject['department']:
                    result += f"    Department: {subject['department']}\n"
            
            result += "\n⚠️ Subjects Needing Attention:\n"
            for subject in reversed(worst_subjects):
                result += f"  - {subject['subject']}: {subject['avg_performance']:.1f}% avg ({subject['attempt_count']} attempts)\n"
                if subject['department']:
                    result += f"    Department: {subject['department']}\n"
            
            # Performance distribution analysis
            cursor.execute(f'''
                SELECT 
                    CASE 
                        WHEN (g.marks * 100.0 / g.max_marks) >= 90 THEN 'A (90-100%)'
                        WHEN (g.marks * 100.0 / g.max_marks) >= 80 THEN 'B (80-89%)'
                        WHEN (g.marks * 100.0 / g.max_marks) >= 70 THEN 'C (70-79%)'
                        WHEN (g.marks * 100.0 / g.max_marks) >= 60 THEN 'D (60-69%)'
                        ELSE 'F (<60%)'
                    END as performance_band,
                    COUNT(*) as count
                FROM grades g
                JOIN students s ON g.student_id = s.id
                WHERE 1=1 {dept_filter}
                GROUP BY performance_band
                ORDER BY 
                    CASE performance_band
                        WHEN 'A (90-100%)' THEN 1
                        WHEN 'B (80-89%)' THEN 2
                        WHEN 'C (70-79%)' THEN 3
                        WHEN 'D (60-69%)' THEN 4
                        WHEN 'F (<60%)' THEN 5
                    END
            ''', params)
            
            distribution = cursor.fetchall()
            if distribution:
                result += "\n📊 Grade Distribution:\n"
                for band in distribution:
                    percentage = (band['count'] / sum(d['count'] for d in distribution)) * 100
                    result += f"  - {band['performance_band']}: {band['count']} records ({percentage:.1f}%)\n"
            
            # Learning recommendations
            result += "\n💡 Learning Insights & Recommendations:\n"
            
            # Calculate overall performance
            overall_avg = sum(s['avg_performance'] * s['attempt_count'] for s in subject_data) / sum(s['attempt_count'] for s in subject_data)
            
            if overall_avg >= 80:
                result += "  ✅ Overall performance is excellent. Maintain current teaching methods.\n"
            elif overall_avg >= 70:
                result += "  📚 Good performance with room for improvement. Consider additional practice materials.\n"
            else:
                result += "  ⚠️ Performance needs attention. Review teaching methods and consider intervention programs.\n"
            
            # Subject-specific recommendations
            if worst_subjects:
                worst_subject = worst_subjects[-1]
                if worst_subject['avg_performance'] < 60:
                    result += f"  🎯 Focus on improving {worst_subject['subject']} - current average: {worst_subject['avg_performance']:.1f}%\n"
            
            return result.strip()
            
        except Exception as e:
            return f"Error analyzing learning trends: {str(e)}"
        finally:
            conn.close()
    def _init_tools(self):
        """Initialize enhanced AI agent tools with analytics"""
        self.tools = [
            StructuredTool.from_function(
                func=self._tool_get_student_info,
                name="get_student_info",
                description="Get detailed information about a specific student by name or ID"
            ),
            StructuredTool.from_function(
                func=self._tool_get_class_statistics,
                name="get_class_statistics",
                description="Get statistics for a department or the entire school"
            ),
            StructuredTool.from_function(
                func=self._tool_get_attendance_summary,
                name="get_attendance_summary",
                description="Get attendance summary for a specific date or overall statistics"
            ),
            StructuredTool.from_function(
                func=self._tool_analyze_student_performance,
                name="analyze_student_performance",
                description="Analyze a student's academic performance and provide insights"
            ),
            StructuredTool.from_function(
                func=self._tool_predict_student_performance,
                name="predict_student_performance",
                description="Predict future student performance using historical data trends"
            ),
            StructuredTool.from_function(
                func=self._tool_analyze_attendance_patterns,
                name="analyze_attendance_patterns",
                description="Analyze attendance patterns and identify at-risk students"
            ),
            StructuredTool.from_function(
                func=self._tool_identify_learning_trends,
                name="identify_learning_trends",
                description="Identify learning trends and subject performance patterns"
            ),
            StructuredTool.from_function(
                func=self.process_uploaded_document,
                name="process_uploaded_document",
                description="Process uploaded documents (PDF, Word, Images) and extract text for analysis"
            ),
            StructuredTool.from_function(
                func=self.search_uploaded_documents,
                name="search_uploaded_documents",
                description="Search within uploaded documents for specific information"
            ),
            StructuredTool.from_function(
                func=self.get_uploaded_documents_summary,
                name="get_uploaded_documents_summary",
                description="Get summary of all uploaded documents"
            )
        ]
    
    def semantic_search(self, query: str, k: int = 5, filters: Dict = None) -> List[Dict]:
        """Enhanced semantic search with similarity scoring"""
        if not self.vector_store:
            return []
        
        try:
            # Create query embedding
            embeddings = OpenAIEmbeddings(api_key=OPENAI_API_KEY)
            query_embedding = embeddings.embed_query(query)
            
            # Perform similarity search with optional filters
            if filters:
                search_kwargs = {"k": k, "filter": filters}
            else:
                search_kwargs = {"k": k}
            
            # Use compression retriever for better results
            if self.compression_retriever:
                docs = self.compression_retriever.get_relevant_documents(query)
            else:
                docs = self.vector_store.similarity_search(query, k=k)
            
            # Calculate similarity scores
            results = []
            for doc in docs:
                # Get document embedding for similarity calculation
                doc_embedding = embeddings.embed_query(doc.page_content)
                
                # Calculate cosine similarity
                similarity = self._cosine_similarity(query_embedding, doc_embedding)
                
                results.append({
                    'content': doc.page_content,
                    'metadata': doc.metadata,
                    'similarity_score': similarity,
                    'relevance': self._classify_relevance(similarity)
                })
            
            # Sort by similarity score
            results.sort(key=lambda x: x['similarity_score'], reverse=True)
            
            return results
            
        except Exception as e:
            print(f"Error in semantic search: {e}")
            return []
    
    def _cosine_similarity(self, vec1: List[float], vec2: List[float]) -> float:
        """Calculate cosine similarity between two vectors"""
        try:
            import numpy as np
            vec1_np = np.array(vec1)
            vec2_np = np.array(vec2)
            
            dot_product = np.dot(vec1_np, vec2_np)
            norm1 = np.linalg.norm(vec1_np)
            norm2 = np.linalg.norm(vec2_np)
            
            if norm1 == 0 or norm2 == 0:
                return 0.0
            
            return float(dot_product / (norm1 * norm2))
        except Exception as e:
            print(f"Error calculating cosine similarity: {e}")
            return 0.0
    
    def _classify_relevance(self, similarity_score: float) -> str:
        """Classify relevance based on similarity score"""
        if similarity_score >= 0.9:
            return "highly_relevant"
        elif similarity_score >= 0.7:
            return "relevant"
        elif similarity_score >= 0.5:
            return "somewhat_relevant"
        else:
            return "not_relevant"
    
    def hybrid_search(self, query: str, k: int = 5) -> List[Dict]:
        """Hybrid search combining semantic and keyword search"""
        results = []
        
        # Semantic search results
        semantic_results = self.semantic_search(query, k)
        
        # Keyword search using TF-IDF
        keyword_results = self._keyword_search(query, k)
        
        # Combine and re-score results
        combined_results = {}
        
        # Add semantic results
        for result in semantic_results:
            doc_id = result['metadata'].get('student_id', result['metadata'].get('chunk_id', 'unknown'))
            combined_results[doc_id] = {
                'content': result['content'],
                'metadata': result['metadata'],
                'semantic_score': result['similarity_score'],
                'keyword_score': 0.0,
                'combined_score': result['similarity_score'] * 0.7  # Weight semantic search higher
            }
        
        # Add keyword results and update scores
        for result in keyword_results:
            doc_id = result['metadata'].get('student_id', result['metadata'].get('chunk_id', 'unknown'))
            if doc_id in combined_results:
                combined_results[doc_id]['keyword_score'] = result['similarity_score']
                combined_results[doc_id]['combined_score'] = (
                    combined_results[doc_id]['semantic_score'] * 0.6 +
                    result['similarity_score'] * 0.4
                )
            else:
                combined_results[doc_id] = {
                    'content': result['content'],
                    'metadata': result['metadata'],
                    'semantic_score': 0.0,
                    'keyword_score': result['similarity_score'],
                    'combined_score': result['similarity_score'] * 0.4
                }
        
        # Sort by combined score
        sorted_results = sorted(
            combined_results.values(),
            key=lambda x: x['combined_score'],
            reverse=True
        )
        
        return sorted_results[:k]
    
    def _keyword_search(self, query: str, k: int = 5) -> List[Dict]:
        """Keyword-based search using TF-IDF"""
        try:
            # Get all documents from vector store
            all_docs = self.vector_store.get()
            
            if not all_docs or not all_docs.get('documents'):
                return []
            
            # Prepare documents and metadata
            documents = all_docs['documents']
            metadatas = all_docs['metadatas']
            
            # Create TF-IDF vectorizer
            vectorizer = TfidfVectorizer(
                stop_words='english',
                max_features=1000,
                ngram_range=(1, 2)
            )
            
            # Fit and transform documents
            tfidf_matrix = vectorizer.fit_transform(documents)
            
            # Transform query
            query_vec = vectorizer.transform([query])
            
            # Calculate cosine similarities
            similarities = cosine_similarity(query_vec, tfidf_matrix).flatten()
            
            # Get top k results
            top_indices = similarities.argsort()[-k:][::-1]
            
            results = []
            for idx in top_indices:
                if similarities[idx] > 0:  # Only include non-zero similarities
                    results.append({
                        'content': documents[idx],
                        'metadata': metadatas[idx] if idx < len(metadatas) else {},
                        'similarity_score': float(similarities[idx]),
                        'relevance': self._classify_relevance(float(similarities[idx]))
                    })
            
            return results
            
        except Exception as e:
            print(f"Error in keyword search: {e}")
            return []
    
            return []
    
    def process_uploaded_document(self, file_path: str, student_id: str = None, description: str = None) -> Dict[str, Any]:
        """Process uploaded documents (PDF, Word, Images) and extract text"""
        try:
            file_path = Path(file_path)
            if not file_path.exists():
                return {"error": "File not found", "success": False}
            
            # Generate file hash for deduplication
            file_hash = self._generate_file_hash(file_path)
            
            # Check if file already processed
            if file_hash in self.uploaded_files:
                return {"message": "File already processed", "file_id": file_hash, "success": True}
            
            # Get file information
            mime_type, _ = mimetypes.guess_type(str(file_path))
            file_size = file_path.stat().st_size
            
            # Extract text based on file type
            extracted_text = ""
            if mime_type == 'application/pdf':
                extracted_text = self._extract_text_from_pdf(file_path)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 'application/msword']:
                extracted_text = self._extract_text_from_docx(file_path)
            elif mime_type and mime_type.startswith('image/'):
                extracted_text = self._extract_text_from_image(file_path)
            elif mime_type and mime_type.startswith('text/'):
                extracted_text = self._extract_text_from_text_file(file_path)
            else:
                return {"error": f"Unsupported file type: {mime_type}", "success": False}
            
            if not extracted_text.strip():
                return {"error": "No text could be extracted from the file", "success": False}
            
            # Create document metadata
            doc_metadata = {
                'file_name': file_path.name,
                'file_path': str(file_path),
                'file_size': file_size,
                'mime_type': mime_type,
                'file_hash': file_hash,
                'student_id': student_id,
                'description': description or f"Document: {file_path.name}",
                'processed_at': datetime.now().isoformat(),
                'type': 'uploaded_document',
                'text_length': len(extracted_text)
            }
            
            # Create document for vector store
            document = Document(
                page_content=extracted_text,
                metadata=doc_metadata
            )
            
            # Add to document store
            self.document_store[file_hash] = document
            self.uploaded_files[file_hash] = doc_metadata
            
            # Update vector store with new document
            self._add_document_to_vector_store(document)
            
            return {
                "success": True,
                "file_id": file_hash,
                "file_name": file_path.name,
                "text_length": len(extracted_text),
                "extracted_preview": extracted_text[:200] + "..." if len(extracted_text) > 200 else extracted_text
            }
            
        except Exception as e:
            return {"error": f"Error processing document: {str(e)}", "success": False}
    
    def _generate_file_hash(self, file_path: Path) -> str:
        """Generate SHA-256 hash for file deduplication"""
        hash_sha256 = hashlib.sha256()
        with open(file_path, "rb") as f:
            for chunk in iter(lambda: f.read(4096), b""):
                hash_sha256.update(chunk)
        return hash_sha256.hexdigest()
    
    def _extract_text_from_pdf(self, file_path: Path) -> str:
        """Extract text from PDF using PyMuPDF"""
        try:
            doc = fitz.open(str(file_path))
            text = ""
            for page_num in range(doc.page_count):
                page = doc.load_page(page_num)
                text += page.get_text() + "\n"
            doc.close()
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from PDF: {e}")
            return ""
    
    def _extract_text_from_docx(self, file_path: Path) -> str:
        """Extract text from Word document"""
        try:
            doc = docx.Document(str(file_path))
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from DOCX: {e}")
            return ""
    
    def _extract_text_from_image(self, file_path: Path) -> str:
        """Extract text from image using OCR"""
        try:
            image = Image.open(file_path)
            # Configure Tesseract for better results
            custom_config = r'--oem 3 --psm 6'
            text = pytesseract.image_to_string(image, config=custom_config)
            return text.strip()
        except Exception as e:
            print(f"Error extracting text from image: {e}")
            return ""
    
    def _extract_text_from_text_file(self, file_path: Path) -> str:
        """Extract text from plain text file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read().strip()
        except UnicodeDecodeError:
            # Try different encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read().strip()
                except UnicodeDecodeError:
                    continue
            return ""
        except Exception as e:
            print(f"Error extracting text from text file: {e}")
            return ""
    
    def _add_document_to_vector_store(self, document: Document):
        """Add processed document to vector store"""
        try:
            if not self.vector_store:
                return
            
            # Split document into chunks
            semantic_splitter = SemanticChunker(
                embeddings=OpenAIEmbeddings(api_key=OPENAI_API_KEY),
                breakpoint_threshold_type="percentile"
            )
            
            try:
                texts = semantic_splitter.split_documents([document])
            except Exception:
                recursive_splitter = RecursiveCharacterTextSplitter(
                    chunk_size=1500,
                    chunk_overlap=300
                )
                texts = recursive_splitter.split_documents([document])
            
            # Add to vector store
            embeddings = OpenAIEmbeddings(api_key=OPENAI_API_KEY)
            self.vector_store.add_documents(texts)
            self.vector_store.persist()
            
            print(f"Added document to vector store: {len(texts)} chunks")
            
        except Exception as e:
            print(f"Error adding document to vector store: {e}")
    
    def search_uploaded_documents(self, query: str, student_id: str = None) -> List[Dict]:
        """Search within uploaded documents"""
        try:
            # Filter documents by student ID if provided
            filters = {'type': 'uploaded_document'}
            if student_id:
                filters['student_id'] = student_id
            
            # Use semantic search
            results = self.semantic_search(query, k=10, filters=filters)
            
            # Enhance results with file information
            enhanced_results = []
            for result in results:
                file_hash = result['metadata'].get('file_hash')
                if file_hash and file_hash in self.uploaded_files:
                    file_info = self.uploaded_files[file_hash]
                    result['file_info'] = file_info
                enhanced_results.append(result)
            
            return enhanced_results
            
        except Exception as e:
            print(f"Error searching uploaded documents: {e}")
            return []
    
    def get_uploaded_documents_summary(self) -> Dict[str, Any]:
        """Get summary of all uploaded documents"""
        try:
            total_docs = len(self.uploaded_files)
            if total_docs == 0:
                return {"total_documents": 0, "message": "No documents uploaded"}
            
            # Analyze document types
            doc_types = {}
            student_docs = {}
            total_size = 0
            
            for file_hash, metadata in self.uploaded_files.items():
                # Count by type
                mime_type = metadata.get('mime_type', 'unknown')
                doc_types[mime_type] = doc_types.get(mime_type, 0) + 1
                
                # Count by student
                student_id = metadata.get('student_id', 'general')
                student_docs[student_id] = student_docs.get(student_id, 0) + 1
                
                total_size += metadata.get('file_size', 0)
            
            return {
                "total_documents": total_docs,
                "total_size_mb": round(total_size / (1024 * 1024), 2),
                "document_types": doc_types,
                "documents_by_student": student_docs,
                "last_updated": datetime.now().isoformat()
            }
            
        except Exception as e:
            return {"error": f"Error getting documents summary: {str(e)}"}
    
    def start_conversation_session(self, session_id: str, user_context: Dict = None) -> Dict[str, Any]:
        """Start a new conversation session with context"""
        try:
            # Initialize session context
            self.session_context[session_id] = {
                'session_id': session_id,
                'started_at': datetime.now().isoformat(),
                'user_context': user_context or {},
                'conversation_history': [],
                'topic_summary': '',
                'last_activity': datetime.now().isoformat(),
                'message_count': 0
            }
            
            # Initialize session-specific memory
            session_memory = ConversationBufferMemory(
                memory_key=f'chat_history_{session_id}',
                return_messages=True,
                output_key='answer'
            )
            
            self.session_context[session_id]['memory'] = session_memory
            
            return {
                "success": True,
                "session_id": session_id,
                "message": "Conversation session started successfully"
            }
            
        except Exception as e:
            return {"error": f"Error starting session: {str(e)}", "success": False}
    
    def add_to_conversation(self, session_id: str, user_message: str, ai_response: str) -> Dict[str, Any]:
        """Add message to conversation history and update context"""
        try:
            if session_id not in self.session_context:
                return {"error": "Session not found", "success": False}
            
            session = self.session_context[session_id]
            
            # Add to conversation history
            session['conversation_history'].append({
                'timestamp': datetime.now().isoformat(),
                'user_message': user_message,
                'ai_response': ai_response,
                'message_type': 'exchange'
            })
            
            # Update memory
            if 'memory' in session and session['memory']:
                session['memory'].chat_memory.add_user_message(user_message)
                session['memory'].chat_memory.add_ai_message(ai_response)
            
            # Update session metadata
            session['last_activity'] = datetime.now().isoformat()
            session['message_count'] += 1
            
            # Generate topic summary every 5 messages
            if session['message_count'] % 5 == 0:
                session['topic_summary'] = self._generate_topic_summary(session_id)
            
            return {
                "success": True,
                "message_count": session['message_count'],
                "topic_summary": session['topic_summary']
            }
            
        except Exception as e:
            return {"error": f"Error adding to conversation: {str(e)}", "success": False}
    
    def _generate_topic_summary(self, session_id: str) -> str:
        """Generate summary of conversation topics"""
        try:
            if session_id not in self.session_context:
                return ""
            
            session = self.session_context[session_id]
            history = session['conversation_history'][-10:]  # Last 10 exchanges
            
            if not history:
                return ""
            
            # Combine recent messages for summary
            conversation_text = "\n".join([
                f"User: {h['user_message']}\nAI: {h['ai_response']}"
                for h in history
            ])
            
            # Generate summary using LLM
            if self.llm:
                prompt = f"""
                Summarize the main topics and themes discussed in this conversation in 2-3 sentences:
                
                {conversation_text}
                
                Focus on: student information, grades, attendance, analytics, or document analysis discussed.
                """
                
                try:
                    summary = self.llm.predict(prompt)
                    return summary.strip()
                except Exception as e:
                    print(f"Error generating summary: {e}")
                    return ""
            
            return ""
            
        except Exception as e:
            print(f"Error in topic summary generation: {e}")
            return ""
    
    def get_conversation_context(self, session_id: str) -> Dict[str, Any]:
        """Get conversation context and history"""
        try:
            if session_id not in self.session_context:
                return {"error": "Session not found", "success": False}
            
            session = self.session_context[session_id]
            
            # Calculate session duration
            started_at = datetime.fromisoformat(session['started_at'])
            last_activity = datetime.fromisoformat(session['last_activity'])
            duration_minutes = (last_activity - started_at).total_seconds() / 60
            
            return {
                "success": True,
                "session_id": session_id,
                "started_at": session['started_at'],
                "last_activity": session['last_activity'],
                "duration_minutes": round(duration_minutes, 2),
                "message_count": session['message_count'],
                "user_context": session['user_context'],
                "topic_summary": session['topic_summary'],
                "recent_history": session['conversation_history'][-5:]  # Last 5 exchanges
            }
            
        except Exception as e:
            return {"error": f"Error getting context: {str(e)}", "success": False}
    
    def end_conversation_session(self, session_id: str) -> Dict[str, Any]:
        """End conversation session and generate final summary"""
        try:
            if session_id not in self.session_context:
                return {"error": "Session not found", "success": False}
            
            session = self.session_context[session_id]
            
            # Generate final summary
            final_summary = self._generate_final_session_summary(session_id)
            
            # Store session summary for future reference
            session_summary = {
                'session_id': session_id,
                'ended_at': datetime.now().isoformat(),
                'duration_minutes': session.get('duration_minutes', 0),
                'message_count': session['message_count'],
                'final_summary': final_summary,
                'user_context': session['user_context']
            }
            
            # Clean up session (optional: keep for analytics)
            # del self.session_context[session_id]
            
            return {
                "success": True,
                "session_summary": session_summary,
                "message": "Session ended successfully"
            }
            
        except Exception as e:
            return {"error": f"Error ending session: {str(e)}", "success": False}
    
    def _generate_final_session_summary(self, session_id: str) -> str:
        """Generate comprehensive summary of the entire session"""
        try:
            if session_id not in self.session_context:
                return ""
            
            session = self.session_context[session_id]
            history = session['conversation_history']
            
            if not history:
                return "No conversation history"
            
            # Analyze conversation patterns
            total_exchanges = len(history)
            topics_discussed = set()
            
            for exchange in history:
                # Simple topic extraction based on keywords
                user_msg = exchange['user_message'].lower()
                if 'grade' in user_msg or 'marks' in user_msg:
                    topics_discussed.add('grades')
                if 'attendance' in user_msg:
                    topics_discussed.add('attendance')
                if 'student' in user_msg:
                    topics_discussed.add('student_info')
                if 'predict' in user_msg or 'trend' in user_msg:
                    topics_discussed.add('analytics')
                if 'document' in user_msg or 'file' in user_msg:
                    topics_discussed.add('documents')
            
            # Generate summary
            if self.llm and total_exchanges > 0:
                conversation_text = "\n".join([
                    f"User: {h['user_message']}\nAI: {h['ai_response']}"
                    for h in history[-10:]  # Last 10 exchanges for context
                ])
                
                prompt = f"""
                Provide a comprehensive summary of this educational consultation session:
                
                Session Details:
                - Total exchanges: {total_exchanges}
                - Topics discussed: {', '.join(topics_discussed) if topics_discussed else 'General conversation'}
                
                Recent conversation:
                {conversation_text}
                
                Generate a 3-4 paragraph summary covering:
                1. Main topics and concerns addressed
                2. Key insights or recommendations provided
                3. Any follow-up actions or areas needing attention
                """
                
                try:
                    summary = self.llm.predict(prompt)
                    return summary.strip()
                except Exception as e:
                    print(f"Error generating final summary: {e}")
                    
            # Fallback summary
            return f"Session included {total_exchanges} exchanges covering topics: {', '.join(topics_discussed) if topics_discussed else 'General conversation'}"
            
        except Exception as e:
            print(f"Error in final summary generation: {e}")
            return "Summary generation failed"
    
    def get_active_sessions(self) -> List[Dict[str, Any]]:
        """Get list of all active sessions"""
        try:
            active_sessions = []
            current_time = datetime.now()
            
            for session_id, session in self.session_context.items():
                last_activity = datetime.fromisoformat(session['last_activity'])
                inactive_minutes = (current_time - last_activity).total_seconds() / 60
                
                active_sessions.append({
                    'session_id': session_id,
                    'message_count': session['message_count'],
                    'last_activity': session['last_activity'],
                    'inactive_minutes': round(inactive_minutes, 2),
                    'topic_summary': session['topic_summary'],
                    'user_context': session['user_context']
                })
            
            # Sort by last activity (most recent first)
            active_sessions.sort(key=lambda x: x['last_activity'], reverse=True)
            
            return active_sessions
            
        except Exception as e:
            print(f"Error getting active sessions: {e}")
            return []
    
            return []
    
    def cache_data(self, key: str, data: Any, ttl_minutes: int = 30) -> bool:
        """Cache data with TTL (time to live)"""
        try:
            import time
            current_time = time.time()
            
            self.cache[key] = data
            self.cache_timestamps[key] = {
                'created_at': current_time,
                'ttl': ttl_minutes * 60,  # Convert to seconds
                'expires_at': current_time + (ttl_minutes * 60)
            }
            
            return True
        except Exception as e:
            print(f"Error caching data: {e}")
            return False
    
    def get_cached_data(self, key: str) -> Any:
        """Retrieve cached data if still valid"""
        try:
            import time
            current_time = time.time()
            
            if key not in self.cache or key not in self.cache_timestamps:
                return None
            
            timestamp_info = self.cache_timestamps[key]
            
            # Check if cache has expired
            if current_time > timestamp_info['expires_at']:
                # Remove expired cache
                del self.cache[key]
                del self.cache_timestamps[key]
                return None
            
            return self.cache[key]
            
        except Exception as e:
            print(f"Error retrieving cached data: {e}")
            return None
    
    def invalidate_cache(self, pattern: str = None) -> int:
        """Invalidate cache entries, optionally by pattern"""
        try:
            keys_to_remove = []
            
            if pattern:
                # Remove keys matching pattern
                import re
                for key in self.cache.keys():
                    if re.match(pattern, key):
                        keys_to_remove.append(key)
            else:
                # Remove all cache entries
                keys_to_remove = list(self.cache.keys())
            
            # Remove from both cache and timestamps
            for key in keys_to_remove:
                self.cache.pop(key, None)
                self.cache_timestamps.pop(key, None)
            
            return len(keys_to_remove)
            
        except Exception as e:
            print(f"Error invalidating cache: {e}")
            return 0
    
    def get_cache_stats(self) -> Dict[str, Any]:
        """Get cache statistics"""
        try:
            import time
            current_time = time.time()
            
            total_entries = len(self.cache)
            expired_entries = 0
            
            for key, timestamp_info in self.cache_timestamps.items():
                if current_time > timestamp_info['expires_at']:
                    expired_entries += 1
            
            return {
                'total_entries': total_entries,
                'expired_entries': expired_entries,
                'valid_entries': total_entries - expired_entries,
                'cache_keys': list(self.cache.keys()),
                'last_cleanup': current_time
            }
            
        except Exception as e:
            return {"error": f"Error getting cache stats: {str(e)}"}
    
    def queue_sync_operation(self, operation_type: str, data: Dict[str, Any]) -> bool:
        """Queue a data synchronization operation"""
        try:
            sync_operation = {
                'operation_type': operation_type,
                'data': data,
                'timestamp': datetime.now().isoformat(),
                'status': 'pending'
            }
            
            self.sync_queue.append(sync_operation)
            
            # Process sync queue if it's getting large
            if len(self.sync_queue) >= 10:
                self.process_sync_queue()
            
            return True
            
        except Exception as e:
            print(f"Error queuing sync operation: {e}")
            return False
    
    def process_sync_queue(self) -> Dict[str, Any]:
        """Process pending synchronization operations"""
        try:
            if not self.sync_queue:
                return {"processed": 0, "message": "No operations to sync"}
            
            processed_count = 0
            failed_operations = []
            
            # Create a copy of the queue to avoid modification during iteration
            operations_to_process = self.sync_queue.copy()
            self.sync_queue.clear()
            
            for operation in operations_to_process:
                try:
                    success = self._execute_sync_operation(operation)
                    
                    if success:
                        processed_count += 1
                        operation['status'] = 'completed'
                    else:
                        failed_operations.append(operation)
                        operation['status'] = 'failed'
                        
                except Exception as e:
                    print(f"Error executing sync operation: {e}")
                    failed_operations.append(operation)
                    operation['status'] = 'failed'
            
            # Update last sync time
            self.last_sync_time = datetime.now().isoformat()
            
            # Re-queue failed operations for retry
            self.sync_queue.extend(failed_operations)
            
            return {
                "processed": processed_count,
                "failed": len(failed_operations),
                "total": len(operations_to_process),
                "last_sync_time": self.last_sync_time
            }
            
        except Exception as e:
            return {"error": f"Error processing sync queue: {str(e)}", "processed": 0}
    
    def _execute_sync_operation(self, operation: Dict[str, Any]) -> bool:
        """Execute a single synchronization operation"""
        try:
            operation_type = operation['operation_type']
            data = operation['data']
            
            if operation_type == 'refresh_vector_store':
                # Refresh vector store with latest data
                self._init_vector_store()
                
            elif operation_type == 'update_student_cache':
                # Update cached student data
                student_id = data.get('student_id')
                if student_id:
                    cache_key = f"student_info_{student_id}"
                    student_info = self._tool_get_student_info(student_id=student_id)
                    self.cache_data(cache_key, student_info, ttl_minutes=15)
                
            elif operation_type == 'update_analytics_cache':
                # Update cached analytics data
                cache_key = "class_statistics"
                stats = self._tool_get_class_statistics()
                self.cache_data(cache_key, stats, ttl_minutes=30)
                
            elif operation_type == 'cleanup_expired_cache':
                # Clean up expired cache entries
                self._cleanup_expired_cache()
                
            else:
                print(f"Unknown sync operation type: {operation_type}")
                return False
            
            return True
            
        except Exception as e:
            print(f"Error executing sync operation: {e}")
            return False
    
    def _cleanup_expired_cache(self) -> int:
        """Clean up expired cache entries"""
        try:
            import time
            current_time = time.time()
            
            expired_keys = []
            for key, timestamp_info in self.cache_timestamps.items():
                if current_time > timestamp_info['expires_at']:
                    expired_keys.append(key)
            
            # Remove expired entries
            for key in expired_keys:
                self.cache.pop(key, None)
                self.cache_timestamps.pop(key, None)
            
            return len(expired_keys)
            
        except Exception as e:
            print(f"Error cleaning up expired cache: {e}")
            return 0
    
    def auto_sync_data(self) -> Dict[str, Any]:
        """Automatically sync data based on configured intervals"""
        try:
            sync_results = {
                'cache_cleanup': 0,
                'vector_store_refreshed': False,
                'analytics_updated': False
            }
            
            # Clean up expired cache
            cleaned_count = self._cleanup_expired_cache()
            sync_results['cache_cleanup'] = cleaned_count
            
            # Check if vector store needs refresh (every hour)
            if self.last_sync_time:
                last_sync = datetime.fromisoformat(self.last_sync_time)
                time_since_sync = (datetime.now() - last_sync).total_seconds() / 60  # minutes
                
                if time_since_sync >= 60:  # 1 hour
                    self.queue_sync_operation('refresh_vector_store', {})
                    sync_results['vector_store_refreshed'] = True
            
            # Update analytics cache (every 30 minutes)
            analytics_cache = self.get_cached_data('class_statistics')
            if not analytics_cache:
                self.queue_sync_operation('update_analytics_cache', {})
                sync_results['analytics_updated'] = True
            
            # Process any pending sync operations
            sync_process_result = self.process_sync_queue()
            sync_results.update(sync_process_result)
            
            return sync_results
            
        except Exception as e:
            return {"error": f"Error in auto sync: {str(e)}"}
    
    def get_sync_status(self) -> Dict[str, Any]:
        """Get current synchronization status"""
        try:
            return {
                'last_sync_time': self.last_sync_time,
                'pending_sync_operations': len(self.sync_queue),
                'cache_stats': self.get_cache_stats(),
                'sync_queue_operations': [
                    {
                        'type': op['operation_type'],
                        'timestamp': op['timestamp'],
                        'status': op['status']
                    }
                    for op in self.sync_queue[-5:]  # Last 5 operations
                ]
            }
            
        except Exception as e:
            return {"error": f"Error getting sync status: {str(e)}"}
    
    def _init_agent(self):
        """Initialize ReAct agent with tools"""
        try:
            if self.llm and self.tools:
                prompt = PromptTemplate.from_template("""
You are an intelligent AI Assistant for the Student Management System.
You help teachers and administrators with student information, grades, attendance, and insights.

You have access to the following tools:
{tools}

Use the following format:

Question: the input question you must answer
Thought: you should always think about what to do
Action: the action to take, should be one of [{tool_names}]
Action Input: the input to the action
Observation: the result of the action
... (this Thought/Action/Action Input/Observation can repeat N times)
Thought: I now know the final answer
Final Answer: the final answer to the original input question

Begin!

Question: {input}
Thought: {agent_scratchpad}
""")
                
                agent = create_react_agent(self.llm, self.tools, prompt)
                self.agent_executor = AgentExecutor(
                    agent=agent,
                    tools=self.tools,
                    verbose=True,
                    handle_parsing_errors=True
                )
        except Exception as e:
            print(f"Error initializing agent: {e}")
            self.agent_executor = None
    
    def query(self, question: str) -> Dict[str, Any]:
        """
        Process a user query and return response
        
        Args:
            question: User's natural language question
            
        Returns:
            Dict with 'answer' and optionally 'sources'
        """
        if not self.llm:
            return {
                'answer': 'AI features are not available. Please check your OpenAI API key configuration.',
                'error': 'LLM not initialized'
            }
        
        try:
            # First try using the agent with tools
            if self.agent_executor:
                response = self.agent_executor.invoke({"input": question})
                return {
                    'answer': response.get('output', 'No response generated'),
                    'sources': []
                }
            
            # Fallback to simple LLM response
            fallback_prompt = f"""
            You are an AI assistant for a Student Management System.
            Answer this question based on your knowledge: {question}
            
            If you need specific student data, mention that the user should provide a student ID or name.
            """
            response = self.llm.predict(fallback_prompt)
            return {
                'answer': response,
                'sources': []
            }
            
        except Exception as e:
            return {
                'answer': f'Error processing your question: {str(e)}',
                'error': str(e)
            }
    
    def get_student_insights(self, student_id: str) -> str:
        """Get AI-generated insights for a specific student"""
        if not self.llm:
            return "AI features not available"
        
        try:
            performance_data = self._tool_analyze_student_performance(student_id)
            
            prompt = f"""
            Based on the following student performance data, provide 3-5 key insights and actionable recommendations:
            
            {performance_data}
            
            Format your response as:
            📊 Key Insights:
            1. ...
            2. ...
            
            💡 Recommendations:
            1. ...
            2. ...
            """
            
            return self.llm.predict(prompt)
            
        except Exception as e:
            return f"Error generating insights: {str(e)}"
    
    def generate_report_summary(self) -> str:
        """Generate a natural language summary of school statistics"""
        if not self.llm:
            return "AI features not available"
        
        try:
            stats = self._tool_get_class_statistics()
            
            prompt = f"""
            As a school administrator, write a brief executive summary (3-4 paragraphs) based on:
            
            {stats}
            
            Include:
            - Overall school health
            - Areas of strength
            - Areas for improvement
            - Actionable next steps
            """
            
            return self.llm.predict(prompt)
            
        except Exception as e:
            return f"Error generating summary: {str(e)}"


# Singleton instance
ai_agent = None

def get_ai_agent():
    """Get or create AI agent singleton"""
    global ai_agent
    global OPENAI_API_KEY
    
    # Reload API key from environment in case it was updated
    refreshed_key = os.getenv('OPENAI_API_KEY')
    if refreshed_key and refreshed_key != OPENAI_API_KEY:
        print(f"API key updated, reinitializing AI agent...")
        OPENAI_API_KEY = refreshed_key
        ai_agent = None  # Force recreation
    
    if ai_agent is None:
        print("Creating new AI agent instance...")
        ai_agent = SMSAIAgent()
    elif not ai_agent.llm:
        print("AI agent exists but LLM not initialized, recreating...")
        ai_agent = SMSAIAgent()
    
    return ai_agent
